import io
import docx
from app.parsers.base import BaseDocumentParser

class DOCXParser(BaseDocumentParser):
    def extract_text(self, file_path: str) -> str:
        """Parse DOCX from a file path (local storage)."""
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs).strip()
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    def extract_text_from_bytes(self, content: bytes) -> str:
        """Parse DOCX directly from bytes (cloud storage — no disk write)."""
        try:
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs).strip()
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX from bytes: {str(e)}")

